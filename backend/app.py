from flask import Flask, request, jsonify, send_file
from PyPDF2 import PdfFileReader, PdfFileWriter
from docx import Document
import cv2
import os
from ebooklib import epub
import tensorflow as tf
from PIL import Image
import io
from kindle_mobi import Mobi
import img2pdf


app = Flask(__name__)

# Reads in a PDF file and returns the text content
def process_pdf(file):
    """Process a PDF file and extract text."""
    reader = PdfFileReader(file)
    text_content = ""
    for page_num in range(reader.getNumPages()):
        page = reader.getPage(page_num)
        text_content += page.extractText()
    return text_content


# Reads in a DOCX file and returns the text content
def process_docx(file):
    """Process a DOCX file and extract text."""
    doc = Document(file)
    text_content = "\n".join([para.text for para in doc.paragraphs])
    return text_content

# Resize an image to the specified dimensions
def resize_image(image_path, width, height):
    """Resize an image to the specified dimensions."""
    img = cv2.imread(image_path)
    resized_img = cv2.resize(img, (width, height))
    resized_path = 'resized_' + os.path.basename(image_path)
    cv2.imwrite(resized_path, resized_img)
    return resized_path

# Create an EPUB file from text content
def create_epub(text_content, title="Sample Ebook", author="Author Name"):
    """Create an EPUB file from text content."""
    book = epub.EpubBook()
    book.set_title(title)
    book.add_author(author)
    chapter = epub.EpubHtml(title='Chapter 1', file_name='chap_01.xhtml', content=text_content)
    book.add_item(chapter)
    epub_path = 'sample.epub'
    epub.write_epub(epub_path, book, {})
    return epub_path

# Convert an EPUB file to MOBI format
def convert_to_mobi(epub_path):
    """Convert an EPUB file to MOBI format."""
    mobi_path = epub_path.replace('.epub', '.mobi')
    Mobi(epub_path).write(mobi_path)
    return mobi_path


# Compress a PDF file
def compress_pdf(file):
    """Compress a PDF file."""
    reader = PdfFileReader(file)
    writer = PdfFileWriter()
    for page_num in range(reader.getNumPages()):
        page = reader.getPage(page_num)
        writer.addPage(page)
    compressed_path = 'compressed.pdf'
    with open(compressed_path, 'wb') as output_file:
        writer.write(output_file)
    return compressed_path


# Apply AI-based image enhancement
def enhance_image(image_path):
    """Apply AI-based image enhancement."""
    img = Image.open(image_path)
    # Apply AI-based enhancement techniques (e.g., super-resolution, denoising)
    # Here, we're just applying a simple sharpening filter as an example
    enhanced_img = img.filter(Image.SHARPEN)
    enhanced_path = 'enhanced_' + os.path.basename(image_path)
    enhanced_img.save(enhanced_path)
    return enhanced_path

# Convert multiple images to a PDF file
def convert_images_to_pdf(image_paths):
    """Convert multiple images to a PDF file."""
    pdf_path = 'converted.pdf'
    with open(pdf_path, 'wb') as f:
        f.write(img2pdf.convert(image_paths))
    return pdf_path

# Convert a PDF file to multiple images
def convert_pdf_to_images(pdf_path):
    """Convert a PDF file to multiple images."""
    images = []
    with open(pdf_path, 'rb') as f:
        pdf = PdfFileReader(f)
        for page_num in range(pdf.getNumPages()):
            page = pdf.getPage(page_num)
            # Extract images from the page and save them
            # Append the image paths to the `images` list
    return images

# Extract metadata from a PDF or DOCX file
def extract_metadata(file):
    """Extract metadata from a PDF or DOCX file."""
    metadata = {}
    if file.filename.endswith('.pdf'):
        reader = PdfFileReader(file)
        metadata = reader.getDocumentInfo()
    elif file.filename.endswith('.docx'):
        doc = Document(file)
        core_properties = doc.core_properties
        metadata = {
            'author': core_properties.author,
            'title': core_properties.title,
            'subject': core_properties.subject,
            'keywords': core_properties.keywords,
        }
    return metadata


# Merge PDFS
def merge_pdfs(pdf_files, output_path='merged.pdf'):
    """Merge multiple PDF files into one."""
    writer = PdfFileWriter()
    for pdf_file in pdf_files:
        reader = PdfFileReader(pdf_file)
        for page_num in range(reader.getNumPages()):
            page = reader.getPage(page_num)
            writer.addPage(page)
    with open(output_path, 'wb') as output_file:
        writer.write(output_file)
    return output_path

# Split 
def split_pdf(pdf_file, output_dir='split_pages'):
    """Split a PDF into separate files for each page."""
    reader = PdfFileReader(pdf_file)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    for page_num in range(reader.getNumPages()):
        writer = PdfFileWriter()
        writer.addPage(reader.getPage(page_num))
        output_path = os.path.join(output_dir, f'page_{page_num + 1}.pdf')
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
    return output_dir


def convert_docx_to_pdf(docx_path, pdf_path='converted.pdf'):
    """Convert a DOCX file to a PDF."""
    # Using a library like docx2pdf or win32com (on Windows)
    import subprocess
    subprocess.run(['docx2pdf', docx_path, pdf_path])
    return pdf_path

def text_to_speech(text_content, output_path='output.mp3'):
    """Convert text content to speech and save as an MP3 file."""
    from gtts import gTTS
    tts = gTTS(text_content)
    tts.save(output_path)
    return output_path

def ocr_image(image_path):
    """Extract text from an image using OCR."""
    import pytesseract
    img = Image.open(image_path)
    text = pytesseract.image_to_string(img)
    return text

def translate_text(text_content, target_language='es'):
    """Translate text content to a different language."""
    from googletrans import Translator
    translator = Translator()
    translation = translator.translate(text_content, dest=target_language)
    return translation.text

def convert_pdf_to_html(pdf_path, output_dir='html_pages'):
    """Convert PDF pages to HTML format."""
    import pdf2htmlEX
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    output_path = os.path.join(output_dir, os.path.basename(pdf_path).replace('.pdf', '.html'))
    subprocess.run(['pdf2htmlEX', pdf_path, output_path])
    return output_path


def encrypt_pdf(pdf_file, password, output_path='encrypted.pdf'):
    """Encrypt a PDF with a password."""
    reader = PdfFileReader(pdf_file)
    writer = PdfFileWriter()
    for page_num in range(reader.getNumPages()):
        page = reader.getPage(page_num)
        writer.addPage(page)
    writer.encrypt(password)
    with open(output_path, 'wb') as output_file:
        writer.write(output_file)
    return output_path

def decrypt_pdf(encrypted_pdf_file, password, output_path='decrypted.pdf'):
    """Decrypt an encrypted PDF with a password."""
    reader = PdfFileReader(encrypted_pdf_file)
    if reader.decrypt(password):
        writer = PdfFileWriter()
        for page_num in range(reader.getNumPages()):
            page = reader.getPage(page_num)
            writer.addPage(page)
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        return output_path
    else:
        raise ValueError("Incorrect password")


def extract_images_from_pdf(pdf_file, output_dir='extracted_images'):
    """Extract and save all images from a PDF file."""
    from PyPDF2 import PdfFileReader
    import fitz  # PyMuPDF

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    pdf_document = fitz.open(pdf_file)
    image_paths = []

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        images = page.get_images(full=True)

        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_path = os.path.join(output_dir, f"page_{page_num + 1}_img_{img_index + 1}.{image_ext}")
            with open(image_path, "wb") as image_file:
                image_file.write(image_bytes)
            image_paths.append(image_path)
    
    return image_paths


def add_toc_to_pdf(pdf_path, toc, output_path='pdf_with_toc.pdf'):
    """Add a table of contents to a PDF."""
    from PyPDF2 import PdfFileReader, PdfFileWriter
    reader = PdfFileReader(pdf_path)
    writer = PdfFileWriter()

    for page_num in range(reader.getNumPages()):
        writer.addPage(reader.getPage(page_num))

    for entry in toc:
        writer.addBookmark(entry['title'], entry['page'])

    with open(output_path, 'wb') as output_file:
        writer.write(output_file)
    return output_path

def convert_pptx_to_pdf(pptx_path, output_path='presentation.pdf'):
    """Convert a PowerPoint presentation to a PDF."""
    from pptx import Presentation

    prs = Presentation(pptx_path)
    # Export slides to images and then convert to PDF
    # Placeholder code: implement actual logic to convert slides to PDF
    return output_path


@app.route('/upload', methods=['POST'])
def upload_file():
    """Upload and process PDF/DOCX files."""
    file = request.files['file']
    file_type = file.filename.split('.')[-1].lower()
    width = int(request.form.get('width', 800))
    height = int(request.form.get('height', 600))
    text_content = ""

    if file_type == 'pdf':
        text_content = process_pdf(file)
        compressed_path = compress_pdf(file)
    elif file_type == 'docx':
        text_content = process_docx(file)

    epub_path = create_epub(text_content)
    mobi_path = convert_to_mobi(epub_path)
    return jsonify({
        'message': 'File processed',
        'epub_path': epub_path,
        'mobi_path': mobi_path,
        'compressed_path': compressed_path
    })

@app.route('/resize_image', methods=['POST'])
def resize_image_route():
    """Resize an image."""
    file = request.files['file']
    book_type = request.form.get('book_type', 'softcover')
    if book_type == 'softcover':
        width = int(request.form.get('softcover_width', 800))
        height = int(request.form.get('softcover_height', 600))
    else:
        width = int(request.form.get('hardcover_width', 1000))
        height = int(request.form.get('hardcover_height', 800))
    file_path = 'uploaded_' + file.filename
    file.save(file_path)
    resized_path = resize_image(file_path, width, height)
    enhanced_path = enhance_image(resized_path)
    return jsonify({
        'message': 'Image resized and enhanced',
        'resized_path': resized_path,
        'enhanced_path': enhanced_path
    })

@app.route('/convert_images_to_pdf', methods=['POST'])
def convert_images_to_pdf_route():
    """Convert multiple images to a PDF file."""
    files = request.files.getlist('files')
    image_paths = []
    for file in files:
        file_path = 'uploaded_' + file.filename
        file.save(file_path)
        image_paths.append(file_path)
    pdf_path = convert_images_to_pdf(image_paths)
    return jsonify({'message': 'Images converted to PDF', 'pdf_path': pdf_path})

@app.route('/convert_pdf_to_images', methods=['POST'])
def convert_pdf_to_images_route():
    """Convert a PDF file to multiple images."""
    file = request.files['file']
    file_path = 'uploaded_' + file.filename
    file.save(file_path)
    image_paths = convert_pdf_to_images(file_path)
    return jsonify({'message': 'PDF converted to images', 'image_paths': image_paths})

@app.route('/download_epub/<path:filename>', methods=['GET'])
def download_epub(filename):
    """Download an EPUB file."""
    return send_file(filename, as_attachment=True)

@app.route('/download_mobi/<path:filename>', methods=['GET'])
def download_mobi(filename):
    """Download a MOBI file."""
    return send_file(filename, as_attachment=True)

@app.route('/download_compressed_pdf/<path:filename>', methods=['GET'])
def download_compressed_pdf(filename):
    """Download a compressed PDF file."""
    return send_file(filename, as_attachment=True)

@app.route('/download_resized_image/<path:filename>', methods=['GET'])
def download_resized_image(filename):
    """Download a resized image."""
    return send_file(filename, as_attachment=True)

@app.route('/download_enhanced_image/<path:filename>', methods=['GET'])
def download_enhanced_image(filename):
    """Download an enhanced image."""
    return send_file(filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)